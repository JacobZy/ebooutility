# -*- coding=utf-8 -*-
import html2textnew
from docx import Document
from docx.shared import Inches
import urllib2
import urllib
import re
import hypermethods
import codecs
import sqlite3
import hypermethods

bookenlist=['Gen.', 'Ex.', 'Lev.', 'Num.', 'Deut.', 'Josh.', 'Judg.', 'Ruth', '1Sam.', '2Sam.', '1Kings', '2Kings', '1Chron.', '2Chron.', 'Ezra', 'Neh.', 'Est.', 'Job', 'Ps.', 'Prov.', 'Eccles.', 'Song', 'Isa.', 'Jer.', 'Lam.', 'Ezek.', 'Dan.', 'Hos.', 'Joel', 'Amos', 'Obad.', 'Jonah', 'Mic.', 'Nah.', 'Hab.', 'Zeph.', 'Hag.', 'Zech.', 'Mal.', 'Matt.', 'Mark', 'Luke', 'John', 'Acts', 'Rom.', '1Cor.', '2Cor.', 'Gal.', 'Eph.', 'Phil.', 'Col.', '1Thess.', '2Thess.', '1Tim.', '2Tim.', 'Titus', 'Philem.', 'Heb.', 'James', '1Pet.', '2Pet.', '1John', '2John', '3John', 'Jude', 'Rev.']

booklist=[u'创世记',u'出埃及记',u'利未记',u'民数记',u'申命记',u'约书亚记',u'士师记',u'路得记',u'撒母耳记上',u'撒母耳记下',u'列王记上',u'列王记下',u'历代志上',u'历代志下',u'以斯拉记',u'尼希米记',u'以斯贴记',u'约伯记',u'诗篇',u'箴言',u'传道书',u'雅歌',u'以赛亚书',u'耶利米书',u'耶利米哀歌',u'以西结书',u'但以理书',u'何西阿书',u'约珥书',u'阿摩斯书',u'俄巴底亚书',u'约拿书',u'弥迦书',u'拿鸿书',u'哈巴谷书',u'西番亚书',u'哈该书',u'撒迦利亚书',u'玛拉基书',u'马太福音',u'马可福音',u'路加福音',u'约翰福音',u'使徒行传',u'罗马书',u'哥林多前书',u'哥林多后书',u'加拉太书',u'以弗所书',u'非利比书',u'歌罗西书',u'贴撒罗尼迦前书',u'贴撒罗尼迦后书',u'提摩太前书',u'提摩太后书',u'提多书',u'非利门书',u'希伯来书',u'雅各书',u'彼得前书',u'彼得后书',u'约翰一书',u'约翰二书',u'约翰三书',u'犹大书',u'启示录']

simplebooklist=[u'创',u'出',u'利',u'民',u'申',u'书',u'士',u'路',u'撒上',u'撒下',u'王上',u'王下',u'代上',u'代下',u'拉',u'尼',u'斯',u'伯',u'诗',u'箴',u'传',u'歌',u'赛',u'耶',u'哀',u'结',u'但',u'何',u'珥',u'摩',u'俄',u'拿',u'弥',u'鸿',u'哈',u'番',u'该',u'亚',u'玛',u'太',u'可',u'路',u'约',u'徒',u'罗',u'林前',u'林后',u'拉',u'弗',u'腓',u'西',u'贴前',u'贴后',u'提前',u'提后',u'多',u'门',u'来',u'雅',u'彼前',u'彼后',u'约一',u'约二',u'约三',u'犹',u'启']

conn=sqlite3.connect('test.twm')
c=conn.cursor()
xreflist=[]

def add_xref(p,bookn,chaptern,versen):
  sql="select tbi,tci,tvi from xrefs_bcv where fbi="+bookn+" and fci="+chaptern+" and fvi="+versen
  c.execute(sql)
  rs=c.fetchall()
  if len(rs)==0:
    return
  hypermethods.add_internal_hyperlink(p,'xref:'+bookn+'.'+chaptern+'.'+versen,'x')
  xreflist.append((bookn,chaptern,versen,rs))

document = Document()

bookfile=codecs.open('cnu.csv','r','utf-8')

bookoldnum='1'
chapteroldnum='1'
verseoldnum='1'

document.add_heading(booklist[int(bookoldnum)-1], 1)
document.add_heading(chapteroldnum, 2)

rstrong=re.compile(r'<W[HG]\d+>')
for line in bookfile.readlines():
  versetotal=line.split(',')
  if len(versetotal)==0:
    print 'falseverse'+line 
    continue
  booknum=versetotal[0]
  chapternum=versetotal[1]
  versenum=versetotal[2]

  if booknum!=bookoldnum:
    document.add_page_break()
    document.add_heading(booklist[int(booknum)-1], 1)
    document.add_heading(chapternum, 2)
    
  if chapternum!=chapteroldnum:
    document.add_heading(chapternum, 2)
    print chapternum
  versenew=rstrong.sub('',versetotal[3])
  p=document.add_paragraph()
  hypermethods.add_bookmark(p,versetotal[2],'bible:'+booknum+'.'+chapternum+'.'+versenum)
  p.add_run('.'+versenew.strip())
  add_xref(p,booknum,chapternum,versenum)

  bookoldnum=booknum
  chapteroldnum=chapternum
  verseoldnum=versenum

document.add_page_break()
for item in xreflist:
  p=document.add_paragraph()
  hypermethods.add_internal_hyperlink_bookmark(p,'bible:'+item[0]+'.'+item[1]+'.'+item[2],simplebooklist[int(item[0])-1]+item[1]+':'+item[2],'xref:'+item[0]+'.'+item[1]+'.'+item[2])
  p.add_run(u'-交叉参考:')
  p=document.add_paragraph()
  for ver in item[3]:
    hypermethods.add_internal_hyperlink(p,'bible:'+str(ver[0])+'.'+str(ver[1])+'.'+str(ver[2]),simplebooklist[ver[0]-1]+str(ver[1])+':'+str(ver[2]))
    p.add_run(';')

document.save('test.docx')
