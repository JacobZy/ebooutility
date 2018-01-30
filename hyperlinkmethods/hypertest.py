# -*- coding=utf-8 -*-
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.run import Run
import docx

bookmarkid =0

def add_bookmark(paragraph, bookmark_text, bookmark_name):
    global bookmarkid
    tag = paragraph._element    
    start = docx.oxml.shared.OxmlElement('w:bookmarkStart')
    start.set(docx.oxml.ns.qn('w:id'), str(bookmarkid))
    start.set(docx.oxml.ns.qn('w:name'), bookmark_name)
    tag.append(start)

    text = docx.oxml.OxmlElement('w:r')
    text.text = bookmark_text
    tag.append(text)

    end = docx.oxml.shared.OxmlElement('w:bookmarkEnd')
    end.set(docx.oxml.ns.qn('w:id'), str(bookmarkid))
    #end.set(docx.oxml.ns.qn('w:name'), bookmark_name)
    tag.append(end)
    bookmarkid=bookmarkid+1

def add_internal_hyperlink(paragraph,url,name):
	tag = paragraph._element 

	hyperlink = docx.oxml.OxmlElement('w:hyperlink')	
	hyperlink.set(qn('w:anchor'), url)
        
        rt = docx.oxml.OxmlElement('w:r')
              
        rPr = docx.oxml.shared.OxmlElement('w:rPr')
        rstyle = docx.oxml.shared.OxmlElement('w:rStyle')
        rstyle.set(docx.oxml.shared.qn('w:val'), 'Internet')        
        rPr.append(rstyle)
        rt.append(rPr)
	rt.text = name
	hyperlink.append(rt)
        
        tag.append(hyperlink)


def add_hyperlink(paragraph, url, text):
	"""
	A function that places a hyperlink within a paragraph object.

	:param paragraph: The paragraph we are adding the hyperlink to.
	:param url: A string containing the required url
	:param text: The text displayed for the url
	:return: The hyperlink object
	"""

	# This gets access to the document.xml.rels file and gets a new relation id value
	part = paragraph.part
	r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

	# Create the w:hyperlink tag and add needed values
	hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
	hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)
        print r_id

	# Create a w:r element
	new_run = docx.oxml.shared.OxmlElement('w:r')

	# Create a new w:rPr element
	rPr = docx.oxml.shared.OxmlElement('w:rPr')

	# Join all the xml elements together add add the required text to the w:r element
	new_run.append(rPr)
	new_run.text = text
	hyperlink.append(new_run)

	paragraph._p.append(hyperlink)

	return hyperlink

def add_internal_hyperlink_bookmark(paragraph,url,text,bookmarkname):
	tag = paragraph._element 

	hyperlink = docx.oxml.OxmlElement('w:hyperlink')	
	hyperlink.set(qn('w:anchor'), url)

	global bookmarkid   
	start = docx.oxml.shared.OxmlElement('w:bookmarkStart')
	start.set(docx.oxml.ns.qn('w:id'), str(bookmarkid))
	start.set(docx.oxml.ns.qn('w:name'), bookmarkname)
	hyperlink.append(start)

	end = docx.oxml.shared.OxmlElement('w:bookmarkEnd')
	end.set(docx.oxml.ns.qn('w:id'), str(bookmarkid))
	#end.set(docx.oxml.ns.qn('w:name'), bookmark_name)
	hyperlink.append(end)
	bookmarkid=bookmarkid+1
        
        rt = docx.oxml.OxmlElement('w:r')
              
        rPr = docx.oxml.shared.OxmlElement('w:rPr')
        rstyle = docx.oxml.shared.OxmlElement('w:rStyle')
        rstyle.set(docx.oxml.shared.qn('w:val'), 'Internet')        
        rPr.append(rstyle)
        rt.append(rPr)
	rt.text = text
	hyperlink.append(rt)
        
        tag.append(hyperlink)
#####################################################################################备份#########################

def add_bookmarkold(paragraph, bookmark_text, bookmark_name):
    run = paragraph.add_run()
    tag = document.element.xpath('//w:r')[-1] #run._r  # for reference the following also works: tag =  document.element.xpath('//w:r')[-1]
    start = docx.oxml.shared.OxmlElement('w:bookmarkStart')
    start.set(docx.oxml.ns.qn('w:id'), '0')
    start.set(docx.oxml.ns.qn('w:name'), bookmark_name)
    tag.append(start)

    text = docx.oxml.OxmlElement('w:r')
    text.text = bookmark_text
    tag.append(text)

    end = docx.oxml.shared.OxmlElement('w:bookmarkEnd')
    end.set(docx.oxml.ns.qn('w:id'), '0')
    end.set(docx.oxml.ns.qn('w:name'), bookmark_name)
    tag.append(end)

import docx
document = docx.Document()
p = document.add_paragraph()
add_hyperlink(p, 'http://www.baidu.com', 'baidu')
add_hyperlink(p, 'http://www.baidu.com', 'test')
add_bookmark(p, '32', 'idtest')

p = document.add_paragraph()
add_bookmark(p, '333', 'idtest12')
p.add_run(' and sometttT')
add_internal_hyperlink_bookmark(p,'idtest','testmark','bookmarkname123')
p.add_run('oooooooppppppssssss')
add_internal_hyperlink(p,'idtest12','testmarkok123')
add_internal_hyperlink(p,'bookmarkname123','again')

document.save('demo_hyperlink.docx')
