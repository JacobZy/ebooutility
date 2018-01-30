# -*- coding=utf-8 -*-
import html2textnew
from docx import Document
from docx.shared import Inches
import urllib2
import urllib
import re
import hypermethods

baseurl=u'http://zhsw.org/123/z/释经/查经资料总汇/'

booklist=['Gen.', 'Ex.', 'Lev.', 'Num.', 'Deut.', 'Josh.', 'Judg.', 'Ruth', '1Sam.', '2Sam.', '1Kings', '2Kings', '1Chron.', '2Chron.', 'Ezra', 'Neh.', 'Est.', 'Job', 'Ps.', 'Prov.', 'Eccles.', 'Song', 'Isa.', 'Jer.', 'Lam.', 'Ezek.', 'Dan.', 'Hos.', 'Joel', 'Amos', 'Obad.', 'Jonah', 'Mic.', 'Nah.', 'Hab.', 'Zeph.', 'Hag.', 'Zech.', 'Mal.', 'Matt.', 'Mark', 'Luke', 'John', 'Acts', 'Rom.', '1Cor.', '2Cor.', 'Gal.', 'Eph.', 'Phil.', 'Col.', '1Thess.', '2Thess.', '1Tim.', '2Tim.', 'Titus', 'Philem.', 'Heb.', 'James', '1Pet.', '2Pet.', '1John', '2John', '3John', 'Jude', 'Rev.']

def getContent(url):
  print url
  response=urllib2.urlopen(urllib.quote(url.encode('utf-8'),'/:'))
  html=response.read().decode('gb2312','ignore')
  return html

def getrealurl(url):
  if url.startswith(u'http'):
    return url
  else:
    return baseurl+url


document = Document()

htmltxt=html2textnew.html2text(getContent(baseurl))
rbook=re.compile(u'\[(.+?)\]\((.+?)\)')
books=rbook.findall(htmltxt)
bookid=0
for item in books:
  booktitle=item[0].strip()
  print booktitle
  
  document.add_heading(booktitle, 1)
  bookurl=getrealurl(item[1].strip())
  print bookurl
  authorstext=html2textnew.html2text(getContent(bookurl))
  print authorstext
  rauthor=re.compile(u'\[马唐纳圣经注释\]\((.+?)\)')
  authortrip=re.compile(r'\n')
  author=rauthor.search(authortrip.sub('',authorstext))
  if author==None:
    bookid=bookid+1
    continue
  authorurl=author.group(1)
  authorurl=re.match(r'(http.+?)index.html',bookurl).group(1)+authorurl
  print authorurl
  chaptercontext=getContent(authorurl)
  chapterstext=html2textnew.html2text(chaptercontext)
  
  print chapterstext

  rchapter=re.compile(u'\)\[(.+?)\]\((.+?)\)')
  chapters=rchapter.findall(chapterstext)

  chapterid=0
  if len(chapters)>0:
    p = document.add_paragraph()
    p.add_run('| ').italic = True
    for chapteritem in chapters:      
      chaptertitle=chapteritem[0].strip()    
      hypermethods.add_internal_hyperlink(p,booklist[bookid]+str(chapterid),chaptertitle)
      p.add_run(' |').italic = True
      chapterid=chapterid+1

  chapterid=0
  for chapteritem in chapters:
    chaptertitle=chapteritem[0].strip()
    print chaptertitle
    
    #document.add_heading(chaptertitle, level=2)
    p = document.add_paragraph()
    hypermethods.add_bookmark(p,chaptertitle,booklist[bookid]+str(chapterid))
    chapterurl=re.match(r'(http.+?)index.html',bookurl).group(1)+chapteritem[1].strip()
    print chapterurl
    chapterstext=html2textnew.html2text(getContent(chapterurl))

    chaptertrip=re.compile(u'\[上一篇\].+?\[下一篇\].+?\)')
    chapterstext=chaptertrip.sub('',chapterstext, 10)
    chaptertrip=re.compile(chaptertitle)
    chapterstext=chaptertrip.sub('',chapterstext, 1)
    chaptertrip=re.compile(r'\{\\Section\:TopicID=\d+\}')
    chapterstext=chaptertrip.sub('',chapterstext)
    chaptertrip=re.compile(u'\*+')
    chapterstext=chaptertrip.sub('',chapterstext)

    #print chapterstext
    p = document.add_paragraph(chapterstext)
    chapterid=chapterid+1

  bookid=bookid+1
  

document.save('test.docx')
