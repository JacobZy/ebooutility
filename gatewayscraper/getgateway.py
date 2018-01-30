# -*- coding=utf-8 -*-
import html2textnew
from docx import Document
from docx.shared import Inches
import urllib2
import urllib
import re
import hypermethods

baseurl=u'https://www.biblegateway.com'
booklisturl=u'https://www.biblegateway.com/versions/Chinese-Union-Version-Simplified-CUVS/#booklist'
chapterurl=u'https://www.biblegateway.com/passage/?search=%E5%87%BA+%E5%9F%83+%E5%8F%8A+%E8%A8%98+3&version=CUVS'

booklist=['Gen.', 'Ex.', 'Lev.', 'Num.', 'Deut.', 'Josh.', 'Judg.', 'Ruth', '1Sam.', '2Sam.', '1Kings', '2Kings', '1Chron.', '2Chron.', 'Ezra', 'Neh.', 'Est.', 'Job', 'Ps.', 'Prov.', 'Eccles.', 'Song', 'Isa.', 'Jer.', 'Lam.', 'Ezek.', 'Dan.', 'Hos.', 'Joel', 'Amos', 'Obad.', 'Jonah', 'Mic.', 'Nah.', 'Hab.', 'Zeph.', 'Hag.', 'Zech.', 'Mal.', 'Matt.', 'Mark', 'Luke', 'John', 'Acts', 'Rom.', '1Cor.', '2Cor.', 'Gal.', 'Eph.', 'Phil.', 'Col.', '1Thess.', '2Thess.', '1Tim.', '2Tim.', 'Titus', 'Philem.', 'Heb.', 'James', '1Pet.', '2Pet.', '1John', '2John', '3John', 'Jude', 'Rev.']

def getContent(url):
  print url
  proxy_handler = urllib2.ProxyHandler({'http': '127.0.0.1:8087'})
  opener = urllib2.build_opener(proxy_handler)
  response=opener.open(url)
  return response

def getrealurl(url):
  if url.startswith(u'http'):
    return url
  else:
    return baseurl+url


document = Document()

booklisthander=getContent(booklisturl)
print booklisthander
rbookname=re.compile(r'book-name"><span class="expand icon-expand"></span>(.+?)<span class="num-chapters collapse in">\d+</span></td>')
rchaptersurl=re.compile(r'<td class="chapters collapse">(<a href.+)$')
rchapurl=re.compile(r'<a href="(.+?)" title="(.+?)">')
rcontent=re.compile(r'<span class="chapternum">\d+(.+?)$')
for line in booklisthander.readlines():
  booko=rbookname.search(line)
  print booko
  if None==booko:
      chapterso=rchaptersurl.search(line)
      print chapterso
      if chapterso==None:
        continue
      chaplist=chapterso.group(1)
      chapters=rchapurl.findall(chaplist)
      for item in chapters:
        chaptitle=item[0].strip()
        print chaptitle
        document.add_heading(u'****'+chaptitle+'*****', 1)
        chapurl=item[1].strip()
        contenthander=getContent(baseurl+chapurl)
        content=contenthander.read()
        texto=rcontent.search(content)
        if texto==None:
          print chaptitle+'None text'
          continue
        htmltext=texto.group(1)
        text='1'+html2textnew.html2text(htmltext)
        document.add_paragraph(text)
  booktitle=booko.group(1)
  print booktitle
  document.add_heading(u'*****'+booktitle+'***', 1)

document.save('test.docx')
